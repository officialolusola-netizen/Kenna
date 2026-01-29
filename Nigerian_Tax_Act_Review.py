"""
Script to generate Nigerian Tax Act Review documents:
1. Word document with comprehensive review
2. Excel document with tax tracker
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from openpyxl import Workbook
from openpyxl.styles import Font, Fill, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from datetime import datetime

# ============================================================================
# PART 1: CREATE WORD DOCUMENT
# ============================================================================

def create_word_document():
    doc = Document()

    # Title
    title = doc.add_heading('Comprehensive Review of the Nigerian Tax Act 2025', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph('Effective January 1, 2026')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.italic = True

    doc.add_paragraph()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Background and Legislative Framework',
        '3. Company Income Tax (CIT)',
        '4. Value Added Tax (VAT)',
        '5. Capital Gains Tax (CGT)',
        '6. Withholding Tax (WHT)',
        '7. Stamp Duties',
        '8. Personal Income Tax (PIT)',
        '9. Development Levy',
        '10. Tertiary Education Tax',
        '11. Tax Exemptions and Incentives',
        '12. Filing Deadlines and Payment Timelines',
        '13. Penalties for Non-Compliance',
        '14. Conclusion'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'On June 26, 2025, President Bola Ahmed Tinubu signed four Tax Reform Bills into law, '
        'marking the most significant overhaul of Nigeria\'s tax framework in decades. These laws include:'
    )

    laws_list = [
        'The Nigeria Tax Act (NTA)',
        'The Nigeria Tax Administration Act (NTAA)',
        'The Nigeria Revenue Service Act (NRSA)',
        'The Joint Revenue Board Act (JRBA)'
    ]
    for law in laws_list:
        doc.add_paragraph(law, style='List Bullet')

    doc.add_paragraph(
        'The NTA consolidates and replaces several major tax laws including the Companies Income Tax Act, '
        'Personal Income Tax Act, Petroleum Profits Tax Act, Value Added Tax Act, Capital Gains Tax Act, '
        'and Stamp Duties Act into a unified, modernized regime. All provisions became effective on January 1, 2026.'
    )

    # 2. Background and Legislative Framework
    doc.add_heading('2. Background and Legislative Framework', level=1)
    doc.add_paragraph(
        'The Nigeria Tax Act 2025 represents a comprehensive reform aimed at:'
    )

    objectives = [
        'Streamlining tax compliance for businesses and individuals',
        'Broadening the tax base to improve revenue generation',
        'Eliminating multiple taxation and administrative inefficiencies',
        'Aligning Nigeria\'s tax system with international best practices',
        'Supporting economic development through targeted incentives'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_paragraph(
        'The previous tax framework consisted of multiple separate laws, each with different administrative '
        'requirements, filing deadlines, and enforcement mechanisms. The NTA consolidates these into a single '
        'unified statute, administered primarily by the Nigeria Revenue Service (NRS), which succeeds the '
        'Federal Inland Revenue Service (FIRS).'
    )

    # 3. Company Income Tax
    doc.add_heading('3. Company Income Tax (CIT)', level=1)

    doc.add_heading('3.1 Overview', level=2)
    doc.add_paragraph(
        'Company Income Tax is imposed on the profits of all companies operating in Nigeria, whether incorporated '
        'in Nigeria or earning income from Nigerian sources. The tax applies to all profits, gains, and income '
        'from business activities.'
    )

    doc.add_heading('3.2 Tax Rates', level=2)

    # CIT Rates Table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    headers = ['Company Category', 'Annual Turnover', 'CIT Rate']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    cit_data = [
        ['Small Companies', '≤ ₦100 million', '0% (Exempt)'],
        ['Medium Companies', '> ₦100 million - ₦500 million', '20%'],
        ['Large Companies', '> ₦500 million', '30%']
    ]

    for row_idx, row_data in enumerate(cit_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('3.3 Small Company Definition', level=2)
    doc.add_paragraph(
        'A company qualifies as "small" if it meets BOTH of the following criteria:'
    )
    small_criteria = [
        'Annual gross turnover of ₦100 million or less (increased from the previous ₦25 million threshold)',
        'Total fixed assets not exceeding ₦250 million'
    ]
    for criterion in small_criteria:
        doc.add_paragraph(criterion, style='List Bullet')

    doc.add_heading('3.4 Taxable Income', level=2)
    doc.add_paragraph('The following types of income are subject to CIT:')

    taxable_income = [
        'Profits from trade or business activities',
        'Income from professional services',
        'Rental income from property',
        'Dividends received from investments',
        'Interest income',
        'Royalties and licensing fees',
        'Capital gains (now taxed at CIT rates)'
    ]
    for income in taxable_income:
        doc.add_paragraph(income, style='List Bullet')

    doc.add_heading('3.5 Deductible Expenses', level=2)
    doc.add_paragraph('Companies can deduct the following expenses when computing taxable profits:')

    deductible = [
        'Expenses wholly, exclusively, reasonably, and necessarily incurred for the purpose of generating income',
        'Interest expenses on loans used for business purposes (subject to 30% of EBITDA limitation for intercompany loans)',
        'Depreciation/capital allowances on qualifying capital expenditure',
        'Contributions to approved pension schemes',
        'Bad debts (specific provisions only)',
        'Research and development costs'
    ]
    for item in deductible:
        doc.add_paragraph(item, style='List Bullet')

    # 4. Value Added Tax
    doc.add_heading('4. Value Added Tax (VAT)', level=1)

    doc.add_heading('4.1 Overview', level=2)
    doc.add_paragraph(
        'VAT is a consumption tax levied on the supply of goods and services in Nigeria. '
        'It is collected by suppliers and remitted to the tax authority.'
    )

    doc.add_heading('4.2 VAT Rate', level=2)
    doc.add_paragraph(
        'The standard VAT rate remains at 7.5%. Despite early discussions about a potential rate increase, '
        'the NTA maintained the existing rate to support economic stability.'
    )

    doc.add_heading('4.3 VAT Categories', level=2)

    # VAT Categories Table
    vat_table = doc.add_table(rows=4, cols=3)
    vat_table.style = 'Table Grid'

    vat_headers = ['Category', 'Rate', 'Items/Services']
    for i, header in enumerate(vat_headers):
        vat_table.rows[0].cells[i].text = header
        vat_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    vat_data = [
        ['Standard Rate', '7.5%', 'Most goods and services not otherwise exempted or zero-rated'],
        ['Zero-Rated', '0%', 'Basic food items, medical products, pharmaceuticals, educational materials, medical services, education services'],
        ['Exempt', 'N/A', 'Exported goods, diplomatic purchases, rent for residential accommodation']
    ]

    for row_idx, row_data in enumerate(vat_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            vat_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('4.4 Key VAT Changes Under NTA', level=2)
    changes = [
        'Input VAT Recovery: Businesses can now claim input VAT on both services and capital assets (fixed assets), a major change from previous laws that only allowed recovery for goods/materials used for production or resale',
        'E-Invoicing Requirement: All VAT-registered businesses must adopt FIRS-sanctioned e-invoicing systems',
        'Digital Platform Obligations: Digital platforms earning commissions must collect VAT on the full supply value',
        'Non-Resident Registration: Non-resident persons supplying taxable goods/services in Nigeria must register for VAT'
    ]
    for change in changes:
        doc.add_paragraph(change, style='List Bullet')

    doc.add_heading('4.5 VAT-Exempt/Zero-Rated Items (Detailed)', level=2)
    exempt_items = [
        'Basic food items (rice, flour, cooking oil, cereals, livestock produce, staple foods)',
        'All medical and pharmaceutical products (excluding cosmetology and fitness devices)',
        'Medical services',
        'Tuition for nursery through tertiary education',
        'Educational books and materials',
        'Medical equipment',
        'Animal feeds and fertilizers',
        'Exported goods and services'
    ]
    for item in exempt_items:
        doc.add_paragraph(item, style='List Bullet')

    # 5. Capital Gains Tax
    doc.add_heading('5. Capital Gains Tax (CGT)', level=1)

    doc.add_heading('5.1 Overview', level=2)
    doc.add_paragraph(
        'Capital Gains Tax is imposed on gains arising from the disposal of chargeable assets.'
    )

    doc.add_heading('5.2 CGT Rates', level=2)

    cgt_table = doc.add_table(rows=3, cols=2)
    cgt_table.style = 'Table Grid'

    cgt_headers = ['Taxpayer Category', 'CGT Rate']
    for i, header in enumerate(cgt_headers):
        cgt_table.rows[0].cells[i].text = header
        cgt_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    cgt_data = [
        ['Companies', '30% (increased from 10%)'],
        ['Individuals', 'Personal income tax rates apply (progressive rates up to 25%)']
    ]

    for row_idx, row_data in enumerate(cgt_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            cgt_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('5.3 Indirect Offshore Transfers', level=2)
    doc.add_paragraph(
        'A significant change under the NTA is the introduction of CGT on indirect transfers of shares in '
        'Nigerian companies. This means that where shares are sold in offshore holding companies that own '
        'Nigerian entities, CGT is triggered (subject to applicable tax treaty exemptions).'
    )

    doc.add_heading('5.4 CGT Exemptions', level=2)
    exemptions = [
        'Disposals where consideration is less than ₦150 million (increased from ₦100 million) AND chargeable gains do not exceed ₦10 million in any 12 consecutive months',
        'Gains from disposal of shares in licensed startups held for at least 24 months',
        'Small companies are exempt from CGT',
        'Proceeds reinvested in similar assets within one year'
    ]
    for exemption in exemptions:
        doc.add_paragraph(exemption, style='List Bullet')

    # 6. Withholding Tax
    doc.add_heading('6. Withholding Tax (WHT)', level=1)

    doc.add_heading('6.1 Overview', level=2)
    doc.add_paragraph(
        'Withholding Tax is an advance payment of income tax deducted at source from certain payments. '
        'It is not a separate tax but a mechanism for collecting taxes.'
    )

    doc.add_heading('6.2 WHT Rates', level=2)

    wht_table = doc.add_table(rows=12, cols=3)
    wht_table.style = 'Table Grid'

    wht_headers = ['Type of Payment', 'Resident Rate', 'Non-Resident Rate']
    for i, header in enumerate(wht_headers):
        wht_table.rows[0].cells[i].text = header
        wht_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    wht_data = [
        ['Dividends', '10%', '10%'],
        ['Interest', '10%', '10%'],
        ['Royalties', '10%', '10%'],
        ['Rent (on property)', '10%', '10%'],
        ['Commission', '10%', '10%'],
        ['Consultancy/Professional fees', '10%', '10%'],
        ['Technical/Management fees', '10%', '10%'],
        ['Construction contracts', '5%', '5%'],
        ['Supply contracts', '5%', '5%'],
        ['Directors\' fees', '10%', '10%'],
        ['Payments to startups (technical services)', '5%', '5%']
    ]

    for row_idx, row_data in enumerate(wht_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            wht_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('6.3 WHT Exemptions for Small Companies', level=2)
    doc.add_paragraph(
        'Small companies are exempt from deducting tax at source if:'
    )
    wht_exempt = [
        'The transaction value is ₦2 million or less',
        'The supplier has a valid Tax Identification Number (TIN)'
    ]
    for item in wht_exempt:
        doc.add_paragraph(item, style='List Bullet')

    # 7. Stamp Duties
    doc.add_heading('7. Stamp Duties', level=1)

    doc.add_heading('7.1 Overview', level=2)
    doc.add_paragraph(
        'Stamp duty is a tax on legal documents and certain transactions. The NTA integrates all stamp duty '
        'provisions into the unified statute, repealing the standalone Stamp Duties Act.'
    )

    doc.add_heading('7.2 Stamp Duty Rates', level=2)

    stamp_table = doc.add_table(rows=11, cols=2)
    stamp_table.style = 'Table Grid'

    stamp_headers = ['Transaction/Document Type', 'Rate']
    for i, header in enumerate(stamp_headers):
        stamp_table.rows[0].cells[i].text = header
        stamp_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    stamp_data = [
        ['Electronic money transfers (≥₦10,000)', '₦50 per transfer (paid by sender)'],
        ['Leases (< 7 years)', '0.78%'],
        ['Leases (≥ 7 years)', '3%'],
        ['Transfer of mineral assets', '2%'],
        ['Ordinary loan capital', '0.125%'],
        ['Loan capital conversion/consolidation', '0.1%'],
        ['Share capital', '0.75%'],
        ['Agreements and contracts', '₦1,000 flat rate'],
        ['Mortgages (> ₦10 million)', 'Applicable rate'],
        ['Property transfers (> ₦10 million)', 'Applicable rate']
    ]

    for row_idx, row_data in enumerate(stamp_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            stamp_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('7.3 Stamp Duty Exemptions', level=2)
    stamp_exempt = [
        'Electronic transfers below ₦10,000',
        'Salary payments',
        'Intrabank transfers',
        'Transfer of government securities',
        'Transfer of shares and stocks',
        'Low-rent leases (< ₦10 million or 10 times minimum wage)',
        'Property transactions below ₦10 million (mortgages, transfers)'
    ]
    for item in stamp_exempt:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.4 Timing Requirement', level=2)
    doc.add_paragraph('All instruments must be stamped within 30 days of execution.')

    # 8. Personal Income Tax
    doc.add_heading('8. Personal Income Tax (PIT)', level=1)

    doc.add_heading('8.1 Overview', level=2)
    doc.add_paragraph(
        'Personal Income Tax applies to the income of individuals, including employment income, '
        'business income, and investment income.'
    )

    doc.add_heading('8.2 PIT Rates', level=2)

    pit_table = doc.add_table(rows=8, cols=2)
    pit_table.style = 'Table Grid'

    pit_headers = ['Annual Income Band', 'Tax Rate']
    for i, header in enumerate(pit_headers):
        pit_table.rows[0].cells[i].text = header
        pit_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    pit_data = [
        ['Up to ₦800,000', '0% (Exempt)'],
        ['₦800,001 - ₦2,500,000', '15%'],
        ['₦2,500,001 - ₦5,000,000', '17.5%'],
        ['₦5,000,001 - ₦10,000,000', '20%'],
        ['₦10,000,001 - ₦20,000,000', '22.5%'],
        ['Above ₦20,000,000', '25%'],
        ['Minimum Wage Earners', '0% (Fully Exempt)']
    ]

    for row_idx, row_data in enumerate(pit_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            pit_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('8.3 Key PIT Changes', level=2)
    pit_changes = [
        'Elimination of Consolidated Relief Allowance (CRA)',
        'New rent relief: 20% of annual rent paid, subject to maximum of ₦500,000',
        'Increased exemption threshold for loss of employment compensation: ₦10 million to ₦50 million',
        'Exemption for individuals earning minimum wage or ₦800,000 or less annually'
    ]
    for change in pit_changes:
        doc.add_paragraph(change, style='List Bullet')

    # 9. Development Levy
    doc.add_heading('9. Development Levy', level=1)

    doc.add_paragraph(
        'The Development Levy is a new 4% levy on the assessable profits of companies. It replaces '
        'various earmarked taxes including:'
    )

    replaced_taxes = [
        'Tertiary Education Tax',
        'Nigeria Police Trust Fund levy',
        'National Information Technology Development Agency levy',
        'Other earmarked contributions'
    ]
    for tax in replaced_taxes:
        doc.add_paragraph(tax, style='List Bullet')

    doc.add_heading('9.1 Exemptions', level=2)
    levy_exempt = [
        'Small companies (turnover ≤ ₦100 million)',
        'Non-resident companies'
    ]
    for item in levy_exempt:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('9.2 Total Tax Burden', level=2)
    doc.add_paragraph(
        'For large companies, the combined tax burden on profits is approximately 34% '
        '(30% CIT + 4% Development Levy).'
    )

    # 10. Tertiary Education Tax
    doc.add_heading('10. Tertiary Education Tax', level=1)
    doc.add_paragraph(
        'Tertiary Education Tax is imposed on every Nigerian company at 3% of assessable profit '
        'for each year of assessment. This tax is payable within two months of an assessment notice '
        'from the tax authority.'
    )

    doc.add_paragraph(
        'Note: Under the Development Levy framework, this tax is consolidated into the 4% Development Levy.'
    )

    # 11. Tax Exemptions and Incentives
    doc.add_heading('11. Tax Exemptions and Incentives', level=1)

    doc.add_heading('11.1 Small Company Exemptions', level=2)
    small_exempt = [
        'Companies Income Tax (CIT)',
        'Capital Gains Tax (CGT)',
        'Development Levy (4%)',
        'Withholding tax on income and payments to suppliers (under certain conditions)'
    ]
    for item in small_exempt:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('11.2 Economic Development Incentive (EDI)', level=2)
    doc.add_paragraph(
        'The EDI replaces the previous "Pioneer Status" tax holiday with a structured, performance-based incentive:'
    )
    edi_features = [
        '5% annual tax credit for up to 5 years on qualifying capital expenditure',
        'Capital expenditure must be made within 5 years from production date',
        'Unused tax credits can be carried forward for another 5 years'
    ]
    for feature in edi_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('11.3 Sector-Specific Incentives', level=2)
    sector_incentives = [
        'Agriculture: 5-year tax holiday for crop production, livestock, and dairy businesses',
        'Startups: CGT exemption on gains from disposal of shares held for at least 24 months',
        'Venture Capital: Tax exemption on gains from investments in labeled startups through VC funds, accelerators, or incubators',
        'Manufacturing: Zero-rating for essential goods production (allows input VAT recovery)'
    ]
    for incentive in sector_incentives:
        doc.add_paragraph(incentive, style='List Bullet')

    doc.add_heading('11.4 Free Zone Benefits', level=2)
    doc.add_paragraph(
        'Free zones retain tax exemptions for exports up to 25% of sales to the customs territory. '
        'However, this 25% exemption will no longer apply from January 1, 2028, subject to Presidential extension.'
    )

    doc.add_heading('11.5 Global Minimum Tax', level=2)
    doc.add_paragraph(
        'Large companies that are part of multinational groups with global revenue ≥ €750 million or '
        'Nigerian turnover ≥ ₦50 billion must maintain a minimum effective tax rate of 15%. '
        'A top-up tax applies if reliefs or incentives reduce payable tax below 15%.'
    )

    # 12. Filing Deadlines
    doc.add_heading('12. Filing Deadlines and Payment Timelines', level=1)

    deadline_table = doc.add_table(rows=9, cols=3)
    deadline_table.style = 'Table Grid'

    deadline_headers = ['Tax Type', 'Filing/Payment Deadline', 'Notes']
    for i, header in enumerate(deadline_headers):
        deadline_table.rows[0].cells[i].text = header
        deadline_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    deadline_data = [
        ['Company Income Tax (CIT)', '6 months after fiscal year end', 'e.g., June 30 for calendar year companies'],
        ['VAT Returns', '21st of the following month', 'Monthly filing requirement'],
        ['PAYE (Employment Tax)', '10th of the following month', 'Monthly remittance'],
        ['Annual PAYE Returns', 'January 31', 'Annual reconciliation'],
        ['Withholding Tax', '21st of the following month', 'Monthly remittance'],
        ['Capital Gains Tax', 'Within 30 days of disposal', 'Disposal-triggered'],
        ['Tertiary Education Tax', 'Within 2 months of assessment', 'Following CIT assessment'],
        ['Stamp Duty', 'Within 30 days of execution', 'Document-triggered']
    ]

    for row_idx, row_data in enumerate(deadline_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            deadline_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('12.1 Installment Payment Option', level=2)
    doc.add_paragraph(
        'Companies filing within 6 months of their accounting year-end can apply to the NRS '
        'to pay CIT in up to three installments, easing cash flow pressures.'
    )

    # 13. Penalties
    doc.add_heading('13. Penalties for Non-Compliance', level=1)

    penalty_table = doc.add_table(rows=7, cols=2)
    penalty_table.style = 'Table Grid'

    penalty_headers = ['Violation', 'Penalty']
    for i, header in enumerate(penalty_headers):
        penalty_table.rows[0].cells[i].text = header
        penalty_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    penalty_data = [
        ['Late return filing', '₦100,000 for first month + ₦50,000 for each additional month'],
        ['Failure to deduct/remit WHT', '10% of amount not deducted/remitted'],
        ['Late payment of tax', 'Interest at prevailing CBN MPR'],
        ['Awarding contracts to unregistered entities', '₦5 million fine'],
        ['Under-deduction of PAYE', 'Stiffer penalties + interest'],
        ['Failure to maintain proper records', 'Administrative penalties apply']
    ]

    for row_idx, row_data in enumerate(penalty_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            penalty_table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    # 14. Conclusion
    doc.add_heading('14. Conclusion', level=1)
    doc.add_paragraph(
        'The Nigeria Tax Act 2025 represents a landmark reform of Nigeria\'s tax system. Key takeaways for companies include:'
    )

    conclusions = [
        'Small companies (turnover ≤ ₦100 million) enjoy significant exemptions from CIT, CGT, and Development Levy',
        'Medium and large companies face a combined tax burden of approximately 34% (30% CIT + 4% Development Levy)',
        'The CGT rate for companies has tripled from 10% to 30%',
        'VAT remains at 7.5% with expanded input VAT recovery rights',
        'Digital compliance (e-invoicing, electronic recordkeeping) is now mandatory',
        'Pioneer status tax holidays have been replaced with the Economic Development Incentive',
        'Penalties for non-compliance have been significantly increased',
        'Approximately 90% of daily necessities are now zero-rated or exempt from VAT'
    ]
    for conclusion in conclusions:
        doc.add_paragraph(conclusion, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'Companies are advised to review their tax positions, update their systems for digital compliance, '
        'and consult with tax professionals to ensure full compliance with the new regime.'
    )

    # Document Info
    doc.add_paragraph()
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.add_run('Document prepared: ').bold = True
    info_para.add_run(datetime.now().strftime('%B %d, %Y'))

    disclaimer = doc.add_paragraph()
    disclaimer.add_run('Disclaimer: ').bold = True
    disclaimer.add_run(
        'This document is for informational purposes only and does not constitute legal or tax advice. '
        'Please consult with qualified tax professionals for specific guidance.'
    )

    # Save document
    doc.save('/home/user/Kenna/Nigerian_Tax_Act_2025_Comprehensive_Review.docx')
    print("Word document created successfully: Nigerian_Tax_Act_2025_Comprehensive_Review.docx")


# ============================================================================
# PART 2: CREATE EXCEL TRACKER
# ============================================================================

def create_excel_tracker():
    wb = Workbook()

    # Define styles
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    subheader_fill = PatternFill(start_color="2E75B6", end_color="2E75B6", fill_type="solid")
    subheader_font = Font(bold=True, color="FFFFFF", size=10)
    taxable_fill = PatternFill(start_color="FCE4D6", end_color="FCE4D6", fill_type="solid")
    exempt_fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)

    # =========== SHEET 1: Company Tax Tracker ===========
    ws1 = wb.active
    ws1.title = "Company Tax Tracker"

    # Headers
    headers1 = ['Tax Type', 'Applicable To', 'Tax Rate', 'Due Date', 'Payment Timeline',
                'Amount Payable', 'Date Paid', 'Payment Status', 'Reference Number', 'Notes']

    for col, header in enumerate(headers1, 1):
        cell = ws1.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = center_align

    # Data rows
    tax_data = [
        ['Company Income Tax (CIT)', 'Large Companies (>₦500M turnover)', '30%', 'June 30 (for Dec year-end)', '6 months after fiscal year-end', '', '', 'Pending', '', 'Can be paid in 3 installments'],
        ['Company Income Tax (CIT)', 'Medium Companies (₦100M-₦500M)', '20%', 'June 30 (for Dec year-end)', '6 months after fiscal year-end', '', '', 'Pending', '', ''],
        ['Company Income Tax (CIT)', 'Small Companies (≤₦100M)', '0% (Exempt)', 'N/A', 'N/A', '₦0', 'N/A', 'Exempt', 'N/A', 'Turnover ≤₦100M, Assets ≤₦250M'],
        ['Development Levy', 'All Companies (except small)', '4%', 'Same as CIT', 'Same as CIT', '', '', 'Pending', '', 'Replaces various earmarked taxes'],
        ['Development Levy', 'Small Companies', '0% (Exempt)', 'N/A', 'N/A', '₦0', 'N/A', 'Exempt', '', ''],
        ['Value Added Tax (VAT)', 'All registered businesses', '7.5%', '21st of following month', 'Monthly', '', '', 'Pending', '', 'E-invoicing mandatory'],
        ['Capital Gains Tax (CGT)', 'Companies', '30%', 'Within 30 days of disposal', 'Upon asset disposal', '', '', 'Pending', '', 'Increased from 10%'],
        ['Capital Gains Tax (CGT)', 'Small Companies', '0% (Exempt)', 'N/A', 'N/A', '₦0', 'N/A', 'Exempt', '', ''],
        ['Withholding Tax (WHT)', 'All companies (general)', '5-10%', '21st of following month', 'Monthly', '', '', 'Pending', '', 'Varies by payment type'],
        ['Tertiary Education Tax', 'All Nigerian companies', '3%', '2 months after assessment', 'After CIT assessment', '', '', 'Pending', '', 'Based on assessable profit'],
        ['PAYE (as employer)', 'All employers', 'Progressive (0-25%)', '10th of following month', 'Monthly', '', '', 'Pending', '', 'Employee tax deduction'],
        ['Stamp Duty', 'Applicable transactions', 'Varies', 'Within 30 days of execution', 'Per transaction', '', '', 'Pending', '', 'See Stamp Duty sheet']
    ]

    for row_idx, row_data in enumerate(tax_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws1.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = left_align if col_idx <= 2 else center_align
            if 'Exempt' in str(value):
                cell.fill = exempt_fill

    # Adjust column widths
    col_widths1 = [25, 35, 15, 30, 30, 18, 15, 15, 20, 35]
    for col, width in enumerate(col_widths1, 1):
        ws1.column_dimensions[get_column_letter(col)].width = width

    # =========== SHEET 2: Taxable vs Non-Taxable Items ===========
    ws2 = wb.create_sheet("Taxable vs Non-Taxable")

    headers2 = ['Category', 'Item/Transaction', 'Tax Type', 'Taxable?', 'Tax Rate', 'Exemption Conditions', 'Notes']

    for col, header in enumerate(headers2, 1):
        cell = ws2.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = center_align

    items_data = [
        # Income Items
        ['Income', 'Business profits', 'CIT', 'Yes', '30%/20%/0%', 'Small companies exempt', 'Based on company size'],
        ['Income', 'Capital gains', 'CGT', 'Yes', '30%', 'Small companies exempt; <₦150M & <₦10M gain exempt', ''],
        ['Income', 'Dividends received', 'WHT', 'Yes', '10%', 'N/A', 'Withholding at source'],
        ['Income', 'Interest income', 'WHT', 'Yes', '10%', 'N/A', 'Withholding at source'],
        ['Income', 'Rental income', 'CIT/WHT', 'Yes', '10% WHT', 'N/A', ''],
        ['Income', 'Royalties', 'WHT', 'Yes', '10%', 'N/A', ''],

        # VAT Items - Taxable
        ['VAT', 'Professional services', 'VAT', 'Yes', '7.5%', 'N/A', 'Standard rate'],
        ['VAT', 'Telecommunications services', 'VAT', 'Yes', '7.5%', 'N/A', ''],
        ['VAT', 'Banking services', 'VAT', 'Yes', '7.5%', 'N/A', ''],
        ['VAT', 'Hotel accommodation', 'VAT', 'Yes', '7.5%', 'N/A', ''],
        ['VAT', 'Entertainment', 'VAT', 'Yes', '7.5%', 'N/A', ''],

        # VAT Items - Zero-Rated/Exempt
        ['VAT', 'Basic food items (rice, flour, etc.)', 'VAT', 'Zero-Rated', '0%', 'Essential staple foods', 'Input VAT recoverable'],
        ['VAT', 'Medical/pharmaceutical products', 'VAT', 'Zero-Rated', '0%', 'Excl. cosmetology/fitness', ''],
        ['VAT', 'Medical services', 'VAT', 'Zero-Rated', '0%', 'Healthcare services', ''],
        ['VAT', 'Educational services', 'VAT', 'Zero-Rated', '0%', 'Nursery to tertiary', ''],
        ['VAT', 'Educational materials/books', 'VAT', 'Zero-Rated', '0%', 'Learning materials', ''],
        ['VAT', 'Medical equipment', 'VAT', 'Zero-Rated', '0%', 'Healthcare equipment', ''],
        ['VAT', 'Animal feeds', 'VAT', 'Zero-Rated', '0%', 'Agricultural inputs', ''],
        ['VAT', 'Fertilizers', 'VAT', 'Zero-Rated', '0%', 'Agricultural inputs', ''],
        ['VAT', 'Exported goods/services', 'VAT', 'Exempt', '0%', 'Export transactions', ''],
        ['VAT', 'Residential rent', 'VAT', 'Exempt', 'N/A', 'Residential accommodation', ''],

        # Stamp Duty Items
        ['Stamp Duty', 'Electronic transfers ≥₦10,000', 'Stamp Duty', 'Yes', '₦50 per transfer', 'Transfers <₦10,000 exempt', 'Sender pays'],
        ['Stamp Duty', 'Salary payments', 'Stamp Duty', 'No', 'N/A', 'Fully exempt', ''],
        ['Stamp Duty', 'Intrabank transfers', 'Stamp Duty', 'No', 'N/A', 'Fully exempt', ''],
        ['Stamp Duty', 'Share transfers', 'Stamp Duty', 'No', 'N/A', 'Fully exempt', ''],
        ['Stamp Duty', 'Property lease (<7 years)', 'Stamp Duty', 'Yes', '0.78%', 'Low-rent leases exempt', ''],
        ['Stamp Duty', 'Property lease (≥7 years)', 'Stamp Duty', 'Yes', '3%', 'Low-rent leases exempt', ''],
        ['Stamp Duty', 'Agreements/contracts', 'Stamp Duty', 'Yes', '₦1,000 flat', 'N/A', ''],

        # WHT Items
        ['WHT', 'Consultancy fees', 'WHT', 'Yes', '10%', 'Small company exemption conditions', ''],
        ['WHT', 'Construction contracts', 'WHT', 'Yes', '5%', '', ''],
        ['WHT', 'Supply contracts', 'WHT', 'Yes', '5%', '', ''],
        ['WHT', 'Directors\' fees', 'WHT', 'Yes', '10%', '', ''],
        ['WHT', 'Technical services to startups', 'WHT', 'Yes', '5%', '', 'Reduced rate for startups'],
    ]

    for row_idx, row_data in enumerate(items_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws2.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = left_align if col_idx <= 2 else center_align
            if row_data[3] in ['Zero-Rated', 'Exempt', 'No']:
                cell.fill = exempt_fill
            elif row_data[3] == 'Yes':
                cell.fill = taxable_fill

    col_widths2 = [15, 35, 12, 12, 20, 35, 30]
    for col, width in enumerate(col_widths2, 1):
        ws2.column_dimensions[get_column_letter(col)].width = width

    # =========== SHEET 3: Filing Deadlines Calendar ===========
    ws3 = wb.create_sheet("Filing Deadlines")

    headers3 = ['Tax Type', 'Filing Frequency', 'Due Date', 'Covered Period', 'Filed?', 'Date Filed', 'Notes']

    for col, header in enumerate(headers3, 1):
        cell = ws3.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = center_align

    deadlines_data = [
        ['VAT Returns', 'Monthly', '21st of following month', 'January 2026', 'No', '', 'Due: Feb 21, 2026'],
        ['VAT Returns', 'Monthly', '21st of following month', 'February 2026', 'No', '', 'Due: Mar 21, 2026'],
        ['VAT Returns', 'Monthly', '21st of following month', 'March 2026', 'No', '', 'Due: Apr 21, 2026'],
        ['VAT Returns', 'Monthly', '21st of following month', 'April 2026', 'No', '', 'Due: May 21, 2026'],
        ['VAT Returns', 'Monthly', '21st of following month', 'May 2026', 'No', '', 'Due: Jun 21, 2026'],
        ['VAT Returns', 'Monthly', '21st of following month', 'June 2026', 'No', '', 'Due: Jul 21, 2026'],
        ['WHT Returns', 'Monthly', '21st of following month', 'January 2026', 'No', '', 'Due: Feb 21, 2026'],
        ['WHT Returns', 'Monthly', '21st of following month', 'February 2026', 'No', '', 'Due: Mar 21, 2026'],
        ['PAYE Remittance', 'Monthly', '10th of following month', 'January 2026', 'No', '', 'Due: Feb 10, 2026'],
        ['PAYE Remittance', 'Monthly', '10th of following month', 'February 2026', 'No', '', 'Due: Mar 10, 2026'],
        ['Annual PAYE Returns', 'Annual', 'January 31', 'Year 2025', 'No', '', 'Due: Jan 31, 2026'],
        ['CIT Returns', 'Annual', '6 months after FYE', 'Year 2025', 'No', '', 'Due: Jun 30, 2026 (Dec FYE)'],
        ['Tertiary Education Tax', 'Annual', '2 months after assessment', 'Year 2025', 'No', '', 'After CIT assessment'],
    ]

    for row_idx, row_data in enumerate(deadlines_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws3.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = center_align

    col_widths3 = [25, 15, 25, 20, 10, 15, 30]
    for col, width in enumerate(col_widths3, 1):
        ws3.column_dimensions[get_column_letter(col)].width = width

    # =========== SHEET 4: Tax Rates Summary ===========
    ws4 = wb.create_sheet("Tax Rates Summary")

    headers4 = ['Tax Type', 'Rate', 'Applicable To', 'Basis of Computation', 'Effective Date']

    for col, header in enumerate(headers4, 1):
        cell = ws4.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = center_align

    rates_data = [
        ['CIT - Large Companies', '30%', 'Turnover > ₦500 million', 'Taxable profits', 'Jan 1, 2026'],
        ['CIT - Medium Companies', '20%', 'Turnover ₦100M - ₦500M', 'Taxable profits', 'Jan 1, 2026'],
        ['CIT - Small Companies', '0%', 'Turnover ≤ ₦100M, Assets ≤ ₦250M', 'Exempt', 'Jan 1, 2026'],
        ['Development Levy', '4%', 'All companies (except small)', 'Assessable profits', 'Jan 1, 2026'],
        ['VAT - Standard', '7.5%', 'Taxable goods/services', 'Transaction value', 'Jan 1, 2026'],
        ['VAT - Zero-Rated', '0%', 'Essential items', 'Transaction value', 'Jan 1, 2026'],
        ['CGT - Companies', '30%', 'All companies (except small)', 'Chargeable gains', 'Jan 1, 2026'],
        ['CGT - Individuals', '0-25%', 'Individuals', 'PIT rates apply', 'Jan 1, 2026'],
        ['WHT - Dividends', '10%', 'All recipients', 'Gross payment', 'Jan 1, 2026'],
        ['WHT - Interest', '10%', 'All recipients', 'Gross payment', 'Jan 1, 2026'],
        ['WHT - Royalties', '10%', 'All recipients', 'Gross payment', 'Jan 1, 2026'],
        ['WHT - Construction', '5%', 'Contractors', 'Contract value', 'Jan 1, 2026'],
        ['WHT - Supply', '5%', 'Suppliers', 'Invoice value', 'Jan 1, 2026'],
        ['PIT - Minimum wage', '0%', 'Minimum wage earners', 'Exempt', 'Jan 1, 2026'],
        ['PIT - ≤₦800K', '0%', 'Low-income individuals', 'Exempt', 'Jan 1, 2026'],
        ['PIT - ₦800K-₦2.5M', '15%', 'Individuals', 'Taxable income', 'Jan 1, 2026'],
        ['PIT - ₦2.5M-₦5M', '17.5%', 'Individuals', 'Taxable income', 'Jan 1, 2026'],
        ['PIT - ₦5M-₦10M', '20%', 'Individuals', 'Taxable income', 'Jan 1, 2026'],
        ['PIT - ₦10M-₦20M', '22.5%', 'Individuals', 'Taxable income', 'Jan 1, 2026'],
        ['PIT - >₦20M', '25%', 'Individuals', 'Taxable income', 'Jan 1, 2026'],
        ['Tertiary Education Tax', '3%', 'All Nigerian companies', 'Assessable profits', 'Jan 1, 2026'],
        ['Stamp Duty - E-Transfer', '₦50', 'Transfers ≥₦10,000', 'Per transaction', 'Jan 1, 2026'],
        ['Stamp Duty - Lease (<7yr)', '0.78%', 'Property leases', 'Annual rent value', 'Jan 1, 2026'],
        ['Stamp Duty - Lease (≥7yr)', '3%', 'Property leases', 'Annual rent value', 'Jan 1, 2026'],
        ['Minimum Effective Tax', '15%', 'Large MNE groups', 'Top-up tax if below 15%', 'Jan 1, 2026'],
    ]

    for row_idx, row_data in enumerate(rates_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws4.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = center_align

    col_widths4 = [28, 15, 35, 25, 15]
    for col, width in enumerate(col_widths4, 1):
        ws4.column_dimensions[get_column_letter(col)].width = width

    # =========== SHEET 5: Exemptions & Incentives ===========
    ws5 = wb.create_sheet("Exemptions & Incentives")

    headers5 = ['Category', 'Beneficiary', 'Exemption/Incentive', 'Conditions', 'Duration', 'Notes']

    for col, header in enumerate(headers5, 1):
        cell = ws5.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = center_align

    exemptions_data = [
        ['Small Companies', 'Turnover ≤₦100M', 'CIT Exemption', 'Turnover ≤₦100M, Assets ≤₦250M', 'Permanent', ''],
        ['Small Companies', 'Turnover ≤₦100M', 'CGT Exemption', 'Same as above', 'Permanent', ''],
        ['Small Companies', 'Turnover ≤₦100M', 'Development Levy Exemption', 'Same as above', 'Permanent', ''],
        ['Small Companies', 'Turnover ≤₦100M', 'WHT Exemption (certain)', 'Transactions ≤₦2M with valid TIN', 'Permanent', ''],
        ['Individuals', 'Low-income earners', 'PIT Exemption', 'Income ≤₦800,000/year', 'Permanent', 'Includes minimum wage earners'],
        ['Agriculture', 'Agri businesses', 'Tax Holiday', 'Crop, livestock, dairy production', '5 years', ''],
        ['Startups', 'Licensed startups', 'CGT Exemption', 'Shares held ≥24 months', 'Upon disposal', ''],
        ['Investors', 'VC/Accelerator investors', 'Tax on gains exemption', 'Investment in labeled startups', 'Permanent', ''],
        ['All Companies', 'Eligible companies', 'Economic Development Incentive', '5% tax credit on capital expenditure', '5 years', 'Can carry forward 5 more years'],
        ['Free Zones', 'FZE entities', 'Export tax exemption', 'Up to 25% sales to customs territory', 'Until Jan 1, 2028', 'Subject to Presidential extension'],
        ['Employment', 'Employees', 'Loss of employment compensation', 'Compensation ≤₦50M', 'Per event', 'Increased from ₦10M'],
        ['Manufacturing', 'Essential goods producers', 'Zero-rated VAT', 'Production of essential goods', 'Permanent', 'Allows input VAT recovery'],
        ['CGT', 'All taxpayers', 'CGT Exemption', 'Disposal <₦150M & gains <₦10M', 'Per 12 months', ''],
    ]

    for row_idx, row_data in enumerate(exemptions_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws5.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = center_align
            cell.fill = exempt_fill

    col_widths5 = [18, 25, 30, 35, 15, 30]
    for col, width in enumerate(col_widths5, 1):
        ws5.column_dimensions[get_column_letter(col)].width = width

    # =========== SHEET 6: Penalties ===========
    ws6 = wb.create_sheet("Penalties")

    headers6 = ['Violation Type', 'Penalty', 'Additional Consequences', 'Notes']

    for col, header in enumerate(headers6, 1):
        cell = ws6.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = center_align

    penalties_data = [
        ['Late return filing', '₦100,000 (first month) + ₦50,000 per additional month', 'Possible prosecution', ''],
        ['Failure to deduct WHT', '10% of amount not deducted', 'Recovery from defaulter', ''],
        ['Failure to remit WHT', '10% of amount not remitted', 'Interest charges', ''],
        ['Late payment of tax', 'Interest at CBN MPR', 'Compounding daily', ''],
        ['Awarding contracts to unregistered entities', '₦5,000,000 fine', 'Disqualification from contracts', ''],
        ['Under-deduction of PAYE', 'Enhanced penalties + interest', 'NRS enforcement', ''],
        ['Failure to maintain records', 'Administrative penalties', 'Audit complications', ''],
        ['Failure to register for VAT', 'Penalties apply', 'Back taxes + interest', ''],
        ['Non-compliance with e-invoicing', 'Penalties apply', 'VAT recovery disallowed', 'Mandatory from Jan 1, 2026'],
    ]

    for row_idx, row_data in enumerate(penalties_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws6.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = center_align
            cell.fill = taxable_fill

    col_widths6 = [35, 45, 30, 30]
    for col, width in enumerate(col_widths6, 1):
        ws6.column_dimensions[get_column_letter(col)].width = width

    # Save workbook
    wb.save('/home/user/Kenna/Nigerian_Tax_Tracker_2026.xlsx')
    print("Excel tracker created successfully: Nigerian_Tax_Tracker_2026.xlsx")


if __name__ == "__main__":
    create_word_document()
    create_excel_tracker()
    print("\nBoth documents have been created successfully!")
