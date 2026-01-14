#!/usr/bin/env python3
"""
Script to create Nigerian Tax Act 2025 Review documents:
1. Word document with comprehensive tax review
2. Excel document with tax payment tracker
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from openpyxl import Workbook
from openpyxl.styles import Font, Fill, PatternFill, Alignment, Border, Side
import os

# ============================================================================
# PART 1: CREATE WORD DOCUMENT - COMPREHENSIVE TAX REVIEW
# ============================================================================

def create_word_document():
    """Create comprehensive Nigerian Tax Act 2025 Review Word document."""

    doc = Document()

    # Set document title
    title = doc.add_heading('COMPREHENSIVE REVIEW OF THE NIGERIAN TAX ACT 2025', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle
    subtitle = doc.add_paragraph('A Complete Guide to Tax Compliance for Companies in Nigeria')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add date
    date_para = doc.add_paragraph('Effective Date: January 1, 2026')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # EXECUTIVE SUMMARY
    # -------------------------------------------------------------------------
    doc.add_heading('1. EXECUTIVE SUMMARY', level=1)

    exec_summary = doc.add_paragraph()
    exec_summary.add_run(
        'On June 26, 2025, President Bola Ahmed Tinubu signed the Nigeria Tax Act, 2025 into law, '
        'marking the most significant overhaul of Nigeria\'s tax framework in over three decades. '
        'The Act consolidates and repeals major tax legislation including the Companies Income Tax Act (CITA), '
        'Personal Income Tax Act (PITA), Petroleum Profits Tax Act (PPTA), Value Added Tax Act (VAT Act), '
        'Capital Gains Tax Act (CGT Act), and Stamp Duties Act (SDA).'
    )

    doc.add_paragraph('The key objectives of the reform include:')
    objectives = [
        'Streamlining tax compliance procedures',
        'Broadening the tax base',
        'Boosting government revenue',
        'Supporting economic development',
        'Eliminating overlapping sector-specific taxes',
        'Modernizing tax administration through digitalization'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # -------------------------------------------------------------------------
    # COMPANY INCOME TAX
    # -------------------------------------------------------------------------
    doc.add_heading('2. COMPANY INCOME TAX (CIT)', level=1)

    doc.add_heading('2.1 Tax Rates', level=2)

    # Create CIT rates table
    cit_table = doc.add_table(rows=4, cols=3)
    cit_table.style = 'Table Grid'

    cit_headers = ['Company Category', 'Criteria', 'Tax Rate']
    cit_data = [
        ['Small Companies', 'Turnover ≤ NGN 50 million AND Fixed Assets ≤ NGN 250 million', '0%'],
        ['Medium Companies', 'Turnover > NGN 50 million to NGN 100 million', '20%'],
        ['Large Companies', 'Turnover > NGN 100 million', '30%']
    ]

    # Add headers
    for i, header in enumerate(cit_headers):
        cell = cit_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Add data
    for row_idx, row_data in enumerate(cit_data):
        for col_idx, cell_data in enumerate(row_data):
            cit_table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('2.2 Taxable Income', level=2)
    doc.add_paragraph('Companies are required to pay tax on the following income categories:')

    taxable_income = [
        'Profits or gains from trade, business, or profession',
        'Rent or premium arising from property use',
        'Dividends, interest, royalties, discounts, and fees',
        'Income from digital services and virtual assets',
        'Capital gains from asset disposals (except for small companies)',
        'Income derived from Nigeria by non-resident companies',
        'Undistributed profits of foreign-controlled companies (CFC rules)',
        'Indirect share transfers of Nigerian assets'
    ]
    for item in taxable_income:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.3 Deductible Expenses', level=2)
    doc.add_paragraph('The following expenses are deductible from assessable profits:')

    deductible = [
        'Expenses wholly, exclusively, and necessarily incurred for business purposes',
        'Interest on borrowed capital (with limitations for connected parties)',
        'Bad debts proven to be bad during the accounting period',
        'Research and development expenses (capped at 5% of turnover)',
        'Capital donations to approved institutions (limited to 10% of profit before tax)',
        'Business expenses incurred up to 6 years before business commencement',
        'Depreciation allowances on qualifying capital expenditure'
    ]
    for item in deductible:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.4 Non-Deductible Expenses', level=2)
    non_deductible = [
        'Capital expenditure (claimed as capital allowances instead)',
        'Personal or domestic expenses',
        'Taxes on profits or gains',
        'Dividends or any profit distribution',
        'Provisions for future liabilities',
        'Bribes and illegal payments',
        'Fines and penalties for breach of law'
    ]
    for item in non_deductible:
        doc.add_paragraph(item, style='List Bullet')

    # -------------------------------------------------------------------------
    # VALUE ADDED TAX
    # -------------------------------------------------------------------------
    doc.add_heading('3. VALUE ADDED TAX (VAT)', level=1)

    doc.add_heading('3.1 VAT Rate', level=2)
    doc.add_paragraph(
        'The VAT rate remains at 7.5% for 2026. Despite earlier proposals to increase the rate to 15%, '
        'the current rate has been maintained to ease the transition to the new tax framework.'
    )

    doc.add_heading('3.2 VAT-Exempt Goods and Services', level=2)
    doc.add_paragraph(
        'The following goods and services are exempt from VAT (suppliers cannot charge VAT and cannot recover input VAT):'
    )

    vat_exempt = [
        'Basic foodstuffs (rice, yam, bread, etc.)',
        'Medical and pharmaceutical products',
        'Educational materials and tuition fees',
        'Residential rent',
        'Public transport fares',
        'Agricultural equipment and inputs',
        'Baby products (food, diapers, etc.)',
        'Assistive devices for persons with disabilities',
        'Humanitarian project goods',
        'Supplies to Export Processing Zones and Free Trade Zones'
    ]
    for item in vat_exempt:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.3 Zero-Rated Supplies', level=2)
    doc.add_paragraph(
        'Zero-rated supplies are subject to VAT at 0%, and suppliers can recover input VAT on related purchases:'
    )

    zero_rated = [
        'Exports of goods and services',
        'Basic food items (expanded list)',
        'Medical products',
        'Educational books and materials',
        'Electricity generation and transmission services',
        'Crude oil exports'
    ]
    for item in zero_rated:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.4 Input VAT Recovery', level=2)
    doc.add_paragraph('Key provisions for input VAT recovery:')
    input_vat = [
        'Registered taxpayers can claim input VAT on taxable supplies including fixed assets',
        'Input VAT on mixed supplies may be claimed in proportion to taxable use',
        'Claims must be made no later than 5 years from incurrence',
        'Refunds for excess input VAT must be claimed within 12 months',
        'Nigeria Revenue Service (NRS) must process refund claims within 30 days'
    ]
    for item in input_vat:
        doc.add_paragraph(item, style='List Bullet')

    # -------------------------------------------------------------------------
    # CAPITAL GAINS TAX
    # -------------------------------------------------------------------------
    doc.add_heading('4. CAPITAL GAINS TAX (CGT)', level=1)

    doc.add_heading('4.1 Tax Rates', level=2)

    cgt_table = doc.add_table(rows=4, cols=3)
    cgt_table.style = 'Table Grid'

    cgt_headers = ['Taxpayer Category', 'Threshold', 'Tax Rate']
    cgt_data = [
        ['Companies (General)', 'Asset sales > NGN 150 million OR gains > NGN 10 million', '30%'],
        ['Small Companies', 'Meet small company criteria', '0%'],
        ['Individuals', 'Capital gains from asset disposal', 'Up to 25%']
    ]

    for i, header in enumerate(cgt_headers):
        cell = cgt_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    for row_idx, row_data in enumerate(cgt_data):
        for col_idx, cell_data in enumerate(row_data):
            cgt_table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    doc.add_heading('4.2 CGT Exemptions', level=2)
    exempt_items = [
        'Disposal of private residence (main home)',
        'Low-value chattels',
        'Up to two personal vehicles',
        'Government bond income and gains',
        'Loss-of-employment compensation (up to NGN 50 million)',
        'Asset disposals by angel investors in startups (held for more than 24 months)',
        'Gains reinvested in similar assets within 12 months (rollover relief)'
    ]
    for item in exempt_items:
        doc.add_paragraph(item, style='List Bullet')

    # -------------------------------------------------------------------------
    # DEVELOPMENT LEVY
    # -------------------------------------------------------------------------
    doc.add_heading('5. DEVELOPMENT LEVY', level=1)

    doc.add_paragraph(
        'A new Development Levy of 4% on assessable profits replaces several overlapping sector-specific taxes:'
    )

    replaced_taxes = [
        'Tertiary Education Tax (TET) - formerly 3%',
        'NASENI Levy',
        'IT Levy',
        'Police Trust Fund Levy'
    ]
    for item in replaced_taxes:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: Small companies and non-resident companies are exempt from the Development Levy.'
    )

    # -------------------------------------------------------------------------
    # STAMP DUTIES
    # -------------------------------------------------------------------------
    doc.add_heading('6. STAMP DUTIES', level=1)

    doc.add_paragraph('Key provisions:')
    stamp_duties = [
        'All instruments must be stamped within 30 days of execution',
        'Electronic stamping is now mandatory',
        'Unstamped instruments are inadmissible as evidence in court',
        'Rates vary based on instrument type and value',
        'Electronic transactions are subject to stamp duties'
    ]
    for item in stamp_duties:
        doc.add_paragraph(item, style='List Bullet')

    # -------------------------------------------------------------------------
    # WITHHOLDING TAX
    # -------------------------------------------------------------------------
    doc.add_heading('7. WITHHOLDING TAX (WHT)', level=1)

    wht_table = doc.add_table(rows=9, cols=3)
    wht_table.style = 'Table Grid'

    wht_headers = ['Payment Type', 'Companies', 'Individuals']
    wht_data = [
        ['Dividends', '10%', '10%'],
        ['Interest', '10%', '10%'],
        ['Rent', '10%', '10%'],
        ['Royalties', '10%', '10%'],
        ['Commissions', '10%', '5%'],
        ['Consultancy/Professional Fees', '10%', '5%'],
        ['Technical Fees', '10%', '5%'],
        ['Construction/Building', '5%', '5%']
    ]

    for i, header in enumerate(wht_headers):
        cell = wht_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    for row_idx, row_data in enumerate(wht_data):
        for col_idx, cell_data in enumerate(row_data):
            wht_table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # TAX EXEMPTIONS
    # -------------------------------------------------------------------------
    doc.add_heading('8. TAX EXEMPTIONS AND RELIEFS', level=1)

    doc.add_heading('8.1 Company Exemptions', level=2)
    company_exemptions = [
        'Small companies (turnover ≤ NGN 50m, fixed assets ≤ NGN 250m) - exempt from CIT, CGT, and Development Levy',
        'Agricultural businesses - first 5 years of operation exempt',
        'Public educational institutions - fully exempt',
        'Charitable and religious organizations - exempt on non-commercial income',
        'Free Zone entities with exports exceeding 75% of sales - 100% exemption',
        'Economic Development Incentive (EDI) - 5% tax credit per annum for 5 years on qualifying capital expenditure'
    ]
    for item in company_exemptions:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8.2 Important Exclusion', level=2)
    doc.add_paragraph(
        'Note: Professional service companies are NOT eligible for small company exemptions, '
        'regardless of whether they meet the financial thresholds. This includes law firms, '
        'accounting firms, consulting firms, medical practices, and similar professional entities.'
    )

    # -------------------------------------------------------------------------
    # TAX PAYMENT TIMELINES
    # -------------------------------------------------------------------------
    doc.add_heading('9. TAX PAYMENT TIMELINES', level=1)

    doc.add_heading('9.1 Company Income Tax', level=2)
    cit_timelines = [
        'Annual returns: Within 6 months from end of accounting year',
        'Newly incorporated companies: Within 18 months from incorporation OR 6 months after accounting period end (whichever is earlier)',
        'From 2026: Monthly installment payments of estimated tax (no longer annual lump sum)',
        'Self-assessment returns: Due with payment'
    ]
    for item in cit_timelines:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('9.2 VAT', level=2)
    vat_timelines = [
        'Monthly filing: On or before 21st of the following month',
        'VAT invoice requirements: Must include business registration number and sequential numbering',
        'Input VAT claims: Within 5 years of incurrence',
        'Refund claims: Within 12 months of excess credit'
    ]
    for item in vat_timelines:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('9.3 PAYE (Pay-As-You-Earn)', level=2)
    paye_timelines = [
        'Monthly remittance: On or before 10th of the following month',
        'Annual reconciliation: Within 31 days after year end',
        'Employee tax deduction cards: Must be provided to employees'
    ]
    for item in paye_timelines:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('9.4 Withholding Tax', level=2)
    wht_timelines = [
        'Monthly remittance: Within 21 days after deduction month',
        'Annual returns: Within 30 days after year end',
        'WHT credit notes must be issued to payees'
    ]
    for item in wht_timelines:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('9.5 Stamp Duties', level=2)
    stamp_timelines = [
        'Stamping deadline: Within 30 days of document execution',
        'Electronic stamping: Mandatory for all instruments',
        'Late stamping penalty: Up to 10x the duty payable'
    ]
    for item in stamp_timelines:
        doc.add_paragraph(item, style='List Bullet')

    # -------------------------------------------------------------------------
    # PENALTIES AND INTEREST
    # -------------------------------------------------------------------------
    doc.add_heading('10. PENALTIES AND INTEREST', level=1)

    penalties_table = doc.add_table(rows=7, cols=2)
    penalties_table.style = 'Table Grid'

    penalty_headers = ['Offense', 'Penalty']
    penalty_data = [
        ['Late filing of returns', 'NGN 50,000 per month (first month) + NGN 25,000 per subsequent month'],
        ['Failure to deduct/remit WHT', 'Penalty + Interest at prevailing Central Bank rate'],
        ['Late payment of tax', 'Interest at CBN rate + 2% per annum'],
        ['Failure to register for VAT', 'NGN 50,000 for first month + NGN 25,000 per subsequent month'],
        ['Non-issuance of VAT invoice', '50% of tax not charged'],
        ['Tax evasion', 'Criminal prosecution + fines + imprisonment']
    ]

    for i, header in enumerate(penalty_headers):
        cell = penalties_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    for row_idx, row_data in enumerate(penalty_data):
        for col_idx, cell_data in enumerate(row_data):
            penalties_table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # COMPLIANCE REQUIREMENTS
    # -------------------------------------------------------------------------
    doc.add_heading('11. COMPLIANCE REQUIREMENTS', level=1)

    doc.add_heading('11.1 Registration Requirements', level=2)
    registration = [
        'Tax registration with NRS (formerly FIRS) immediately after CAC incorporation',
        'VAT registration for businesses making taxable supplies',
        'PAYE registration for employers with employees',
        'Partnerships must register within 30 days of formation'
    ]
    for item in registration:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('11.2 Record Keeping', level=2)
    records = [
        'Maintain books of accounts for at least 6 years',
        'Keep all tax records, receipts, and supporting documents',
        'Maintain employee payroll records',
        'Preserve VAT invoices and receipts'
    ]
    for item in records:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('11.3 Digital Compliance', level=2)
    digital = [
        'Mandatory e-filing through TaxPro-Max platform',
        'VAT fiscalization system and e-invoicing via Merchant Buyer Solution (MBS)',
        'Real-time validation of VAT invoices',
        'Electronic stamping for all instruments'
    ]
    for item in digital:
        doc.add_paragraph(item, style='List Bullet')

    # -------------------------------------------------------------------------
    # SPECIAL PROVISIONS
    # -------------------------------------------------------------------------
    doc.add_heading('12. SPECIAL PROVISIONS', level=1)

    doc.add_heading('12.1 Digital Services and Non-Residents', level=2)
    digital_services = [
        'Non-resident persons supplying taxable goods/services in Nigeria must register for VAT',
        'Digital platforms must collect VAT on full supply value',
        'Streaming subscriptions, digital advertisements, and cloud services are VAT-taxable',
        'Minimum tax for non-resident companies: WHT deducted OR 4% flat rate on Nigerian-source income'
    ]
    for item in digital_services:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('12.2 Controlled Foreign Company (CFC) Rules', level=2)
    cfc_rules = [
        'Nigerian parent companies with foreign subsidiaries below 15% effective tax rate are subject to top-up tax',
        'Minimum effective tax rate (ETR) of 15% for MNE groups',
        'Companies with NGN 20 billion+ turnover subject to minimum ETR requirements'
    ]
    for item in cfc_rules:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('12.3 Transfer Pricing', level=2)
    transfer_pricing = [
        'Arm\'s length principle applies to all connected-party transactions',
        'Transfer pricing returns required for applicable transactions',
        'Documentation requirements for intercompany transactions',
        'Advance pricing agreements available'
    ]
    for item in transfer_pricing:
        doc.add_paragraph(item, style='List Bullet')

    # -------------------------------------------------------------------------
    # CONCLUSION
    # -------------------------------------------------------------------------
    doc.add_heading('13. CONCLUSION AND RECOMMENDATIONS', level=1)

    recommendations = [
        'Review your company\'s tax structure and classification (small, medium, or large)',
        'Update your accounting and compliance systems for monthly tax payments',
        'Ensure VAT registration and proper invoice numbering systems',
        'Review connected-party transactions for transfer pricing compliance',
        'Update employee payroll systems for PAYE compliance',
        'Implement e-invoicing and fiscalization systems before the deadline',
        'Train finance and accounting staff on new requirements',
        'Engage tax advisors for complex transactions and planning'
    ]

    doc.add_paragraph('Companies should take the following steps to ensure compliance:')
    for item in recommendations:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('This document was prepared as a general guide and should not be used as a substitute for professional tax advice. Companies are encouraged to consult with qualified tax professionals for specific guidance on their circumstances.')

    # Save document
    doc.save('/home/user/Kenna/Nigerian_Tax_Act_2025_Comprehensive_Review.docx')
    print("Word document created successfully!")


# ============================================================================
# PART 2: CREATE EXCEL DOCUMENT - TAX TRACKER
# ============================================================================

def create_excel_tracker():
    """Create Nigerian Tax Payment Tracker Excel document."""

    wb = Workbook()

    # Define styles
    header_font = Font(bold=True, color='FFFFFF', size=12)
    header_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    subheader_fill = PatternFill(start_color='2E75B6', end_color='2E75B6', fill_type='solid')
    taxable_fill = PatternFill(start_color='E2EFDA', end_color='E2EFDA', fill_type='solid')
    exempt_fill = PatternFill(start_color='FCE4D6', end_color='FCE4D6', fill_type='solid')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)

    # -------------------------------------------------------------------------
    # SHEET 1: TAX SUMMARY
    # -------------------------------------------------------------------------
    ws1 = wb.active
    ws1.title = 'Tax Summary'

    # Headers
    summary_headers = ['Tax Type', 'Rate', 'Who Pays', 'Due Date', 'Payment Frequency', 'Filing Platform']
    for col, header in enumerate(summary_headers, 1):
        cell = ws1.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    # Data
    summary_data = [
        ['Company Income Tax (CIT) - Large', '30%', 'Companies with turnover > NGN 100m', 'Within 6 months of year end', 'Monthly installments', 'TaxPro-Max'],
        ['Company Income Tax (CIT) - Medium', '20%', 'Companies with turnover NGN 50m-100m', 'Within 6 months of year end', 'Monthly installments', 'TaxPro-Max'],
        ['Company Income Tax (CIT) - Small', '0%', 'Companies with turnover ≤ NGN 50m', 'Exempt', 'N/A', 'N/A'],
        ['Value Added Tax (VAT)', '7.5%', 'All registered businesses', '21st of following month', 'Monthly', 'TaxPro-Max'],
        ['Capital Gains Tax (CGT) - Companies', '30%', 'Companies on asset sales > NGN 150m', 'With CIT return', 'Annual', 'TaxPro-Max'],
        ['Capital Gains Tax (CGT) - Individuals', 'Up to 25%', 'Individuals on capital gains', '90 days after disposal', 'Per transaction', 'TaxPro-Max'],
        ['Development Levy', '4%', 'All companies except small/non-resident', 'With CIT return', 'Annual', 'TaxPro-Max'],
        ['PAYE (Employees)', 'Progressive (0-25%)', 'Employers on behalf of employees', '10th of following month', 'Monthly', 'TaxPro-Max'],
        ['Withholding Tax - Dividends', '10%', 'Companies paying dividends', '21 days after deduction', 'Per payment', 'TaxPro-Max'],
        ['Withholding Tax - Interest', '10%', 'Companies paying interest', '21 days after deduction', 'Per payment', 'TaxPro-Max'],
        ['Withholding Tax - Rent', '10%', 'Companies paying rent', '21 days after deduction', 'Per payment', 'TaxPro-Max'],
        ['Withholding Tax - Consultancy', '10%', 'Companies paying consultants', '21 days after deduction', 'Per payment', 'TaxPro-Max'],
        ['Stamp Duties', 'Variable', 'Parties to instruments', 'Within 30 days of execution', 'Per transaction', 'FIRS e-Stamp'],
    ]

    for row_idx, row_data in enumerate(summary_data, 2):
        for col_idx, cell_data in enumerate(row_data, 1):
            cell = ws1.cell(row=row_idx, column=col_idx, value=cell_data)
            cell.border = thin_border
            cell.alignment = left_align if col_idx in [1, 3] else center_align

    # Adjust column widths
    ws1.column_dimensions['A'].width = 35
    ws1.column_dimensions['B'].width = 15
    ws1.column_dimensions['C'].width = 40
    ws1.column_dimensions['D'].width = 30
    ws1.column_dimensions['E'].width = 20
    ws1.column_dimensions['F'].width = 15

    # -------------------------------------------------------------------------
    # SHEET 2: TAXABLE vs NON-TAXABLE ITEMS
    # -------------------------------------------------------------------------
    ws2 = wb.create_sheet('Taxable vs Exempt Items')

    # Headers
    ws2.merge_cells('A1:D1')
    ws2.cell(row=1, column=1, value='TAXABLE ITEMS').font = header_font
    ws2.cell(row=1, column=1).fill = PatternFill(start_color='006400', end_color='006400', fill_type='solid')
    ws2.cell(row=1, column=1).alignment = center_align

    ws2.merge_cells('F1:I1')
    ws2.cell(row=1, column=6, value='EXEMPT/NON-TAXABLE ITEMS').font = header_font
    ws2.cell(row=1, column=6).fill = PatternFill(start_color='8B0000', end_color='8B0000', fill_type='solid')
    ws2.cell(row=1, column=6).alignment = center_align

    taxable_headers = ['Item/Transaction', 'Tax Type', 'Rate', 'Notes']
    for col, header in enumerate(taxable_headers, 1):
        cell = ws2.cell(row=2, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = taxable_fill
        cell.border = thin_border

    exempt_headers = ['Item/Transaction', 'Tax Type', 'Exemption', 'Notes']
    for col, header in enumerate(exempt_headers, 6):
        cell = ws2.cell(row=2, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = exempt_fill
        cell.border = thin_border

    # Taxable items data
    taxable_items = [
        ['Trading profits/business income', 'CIT', '20-30%', 'Based on company size'],
        ['Rental income (companies)', 'CIT', '30%', 'Part of assessable profit'],
        ['Dividends received', 'WHT/CIT', '10%', 'WHT deductible from CIT'],
        ['Interest income', 'WHT/CIT', '10%', 'WHT deductible from CIT'],
        ['Royalty income', 'WHT/CIT', '10%', 'WHT deductible from CIT'],
        ['Capital gains (large companies)', 'CGT', '30%', 'Sales > NGN 150m or gains > NGN 10m'],
        ['Digital services revenue', 'VAT/CIT', '7.5%/30%', 'New taxable category'],
        ['Virtual/crypto asset gains', 'CGT', '30%', 'New taxable category'],
        ['Commission income', 'WHT/CIT', '10%', 'WHT at source'],
        ['Consultancy fees', 'WHT/CIT', '10%', 'WHT at source'],
        ['Construction contracts', 'WHT/CIT', '5%', 'WHT at source'],
        ['Goods and services (general)', 'VAT', '7.5%', 'Standard taxable supplies'],
        ['Imported goods', 'VAT/Customs', '7.5% + duties', 'At point of import'],
        ['Professional services fees', 'VAT', '7.5%', 'All professional services'],
        ['Gratuity payments', 'PAYE', 'Progressive', 'Now taxable'],
        ['Petroleum profits', 'PPT', '65-85%', 'Special rates apply'],
    ]

    for row_idx, row_data in enumerate(taxable_items, 3):
        for col_idx, cell_data in enumerate(row_data, 1):
            cell = ws2.cell(row=row_idx, column=col_idx, value=cell_data)
            cell.border = thin_border
            cell.alignment = left_align

    # Exempt items data
    exempt_items = [
        ['Small company profits', 'CIT', '0%', 'Turnover ≤ NGN 50m'],
        ['Agricultural income (first 5 years)', 'CIT', 'Exempt', 'New businesses only'],
        ['Private residence gains', 'CGT', 'Exempt', 'Main home only'],
        ['Personal vehicles (up to 2)', 'CGT', 'Exempt', 'Individual exemption'],
        ['Government bond income', 'CIT/CGT', 'Exempt', 'Federal bonds'],
        ['Basic foodstuffs', 'VAT', 'Exempt', 'Rice, yam, bread, etc.'],
        ['Medical products', 'VAT', 'Exempt/Zero-rated', 'Pharmaceutical products'],
        ['Educational materials', 'VAT', 'Exempt/Zero-rated', 'Books, school supplies'],
        ['Residential rent', 'VAT', 'Exempt', 'Housing only'],
        ['Public transport fares', 'VAT', 'Exempt', 'Bus, train fares'],
        ['Baby products', 'VAT', 'Exempt', 'Food, diapers, etc.'],
        ['Disability assistive devices', 'VAT', 'Exempt', 'Wheelchairs, etc.'],
        ['Exports', 'VAT', 'Zero-rated', 'Can recover input VAT'],
        ['Charitable organizations', 'CIT', 'Exempt', 'Non-commercial income only'],
        ['Angel investor gains (>24 months)', 'CGT', 'Exempt', 'Startup investments'],
        ['Loss compensation (up to NGN 50m)', 'PAYE', 'Exempt', 'Employment termination'],
    ]

    for row_idx, row_data in enumerate(exempt_items, 3):
        for col_idx, cell_data in enumerate(row_data, 6):
            cell = ws2.cell(row=row_idx, column=col_idx, value=cell_data)
            cell.border = thin_border
            cell.alignment = left_align

    # Adjust column widths
    for col in ['A', 'F']:
        ws2.column_dimensions[col].width = 35
    for col in ['B', 'C', 'D', 'G', 'H', 'I']:
        ws2.column_dimensions[col].width = 18

    # -------------------------------------------------------------------------
    # SHEET 3: PAYMENT TRACKER
    # -------------------------------------------------------------------------
    ws3 = wb.create_sheet('Payment Tracker')

    # Headers
    tracker_headers = [
        'Tax Type', 'Tax Period', 'Amount Due (NGN)', 'Due Date',
        'Date Paid', 'Amount Paid (NGN)', 'Receipt/Reference No.',
        'Payment Status', 'Notes'
    ]

    for col, header in enumerate(tracker_headers, 1):
        cell = ws3.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    # Sample data rows for the tracker
    sample_data = [
        ['Company Income Tax', 'Jan 2026', '', 'Feb 28, 2026', '', '', '', 'Pending', 'Monthly installment'],
        ['Company Income Tax', 'Feb 2026', '', 'Mar 31, 2026', '', '', '', 'Pending', 'Monthly installment'],
        ['Company Income Tax', 'Mar 2026', '', 'Apr 30, 2026', '', '', '', 'Pending', 'Monthly installment'],
        ['VAT', 'Jan 2026', '', 'Feb 21, 2026', '', '', '', 'Pending', 'Monthly filing'],
        ['VAT', 'Feb 2026', '', 'Mar 21, 2026', '', '', '', 'Pending', 'Monthly filing'],
        ['VAT', 'Mar 2026', '', 'Apr 21, 2026', '', '', '', 'Pending', 'Monthly filing'],
        ['PAYE', 'Jan 2026', '', 'Feb 10, 2026', '', '', '', 'Pending', 'Employer remittance'],
        ['PAYE', 'Feb 2026', '', 'Mar 10, 2026', '', '', '', 'Pending', 'Employer remittance'],
        ['PAYE', 'Mar 2026', '', 'Apr 10, 2026', '', '', '', 'Pending', 'Employer remittance'],
        ['WHT - Rent', 'Q1 2026', '', 'Apr 21, 2026', '', '', '', 'Pending', '21 days after quarter'],
        ['WHT - Consultancy', 'Q1 2026', '', 'Apr 21, 2026', '', '', '', 'Pending', '21 days after quarter'],
        ['Development Levy', 'FY 2025', '', 'Jun 30, 2026', '', '', '', 'Pending', 'With CIT return'],
        ['CGT (if applicable)', 'FY 2025', '', 'Jun 30, 2026', '', '', '', 'Pending', 'With CIT return'],
        ['Annual CIT Return', 'FY 2025', '', 'Jun 30, 2026', '', '', '', 'Pending', '6 months after year end'],
        ['Stamp Duties', 'As incurred', '', 'Within 30 days', '', '', '', 'Pending', 'Per transaction'],
    ]

    for row_idx, row_data in enumerate(sample_data, 2):
        for col_idx, cell_data in enumerate(row_data, 1):
            cell = ws3.cell(row=row_idx, column=col_idx, value=cell_data)
            cell.border = thin_border
            cell.alignment = center_align if col_idx > 1 else left_align

    # Add conditional formatting colors for status
    status_colors = {
        'Pending': 'FFEB9C',
        'Paid': 'C6EFCE',
        'Overdue': 'FFC7CE',
        'In Progress': 'BDD7EE'
    }

    # Adjust column widths
    ws3.column_dimensions['A'].width = 25
    ws3.column_dimensions['B'].width = 15
    ws3.column_dimensions['C'].width = 18
    ws3.column_dimensions['D'].width = 18
    ws3.column_dimensions['E'].width = 15
    ws3.column_dimensions['F'].width = 18
    ws3.column_dimensions['G'].width = 22
    ws3.column_dimensions['H'].width = 15
    ws3.column_dimensions['I'].width = 25

    # -------------------------------------------------------------------------
    # SHEET 4: DUE DATES CALENDAR
    # -------------------------------------------------------------------------
    ws4 = wb.create_sheet('Due Dates Calendar 2026')

    # Headers
    calendar_headers = ['Month', 'Tax Type', 'Due Date', 'Description']
    for col, header in enumerate(calendar_headers, 1):
        cell = ws4.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    # Calendar data
    calendar_data = [
        ['January', 'CIT Monthly Installment', 'Jan 31', 'First monthly payment for 2026'],
        ['January', 'VAT (December 2025)', 'Jan 21', 'Previous month VAT filing'],
        ['January', 'PAYE (December 2025)', 'Jan 10', 'Previous month PAYE remittance'],
        ['February', 'CIT Monthly Installment', 'Feb 28', 'February payment'],
        ['February', 'VAT (January)', 'Feb 21', 'Previous month VAT filing'],
        ['February', 'PAYE (January)', 'Feb 10', 'Previous month PAYE remittance'],
        ['March', 'CIT Monthly Installment', 'Mar 31', 'March payment'],
        ['March', 'VAT (February)', 'Mar 21', 'Previous month VAT filing'],
        ['March', 'PAYE (February)', 'Mar 10', 'Previous month PAYE remittance'],
        ['March', 'Personal Income Tax Return', 'Mar 31', 'Employed individuals (Dec Y/E)'],
        ['April', 'CIT Monthly Installment', 'Apr 30', 'April payment'],
        ['April', 'VAT (March)', 'Apr 21', 'Previous month VAT filing'],
        ['April', 'PAYE (March)', 'Apr 10', 'Previous month PAYE remittance'],
        ['April', 'Q1 WHT Returns', 'Apr 21', 'Quarterly WHT filing'],
        ['May', 'CIT Monthly Installment', 'May 31', 'May payment'],
        ['May', 'VAT (April)', 'May 21', 'Previous month VAT filing'],
        ['May', 'PAYE (April)', 'May 10', 'Previous month PAYE remittance'],
        ['June', 'Annual CIT Return (Dec Y/E)', 'Jun 30', 'Companies with December year end'],
        ['June', 'Development Levy', 'Jun 30', 'With CIT return'],
        ['June', 'VAT (May)', 'Jun 21', 'Previous month VAT filing'],
        ['June', 'PAYE (May)', 'Jun 10', 'Previous month PAYE remittance'],
        ['July', 'VAT (June)', 'Jul 21', 'Previous month VAT filing'],
        ['July', 'PAYE (June)', 'Jul 10', 'Previous month PAYE remittance'],
        ['July', 'Q2 WHT Returns', 'Jul 21', 'Quarterly WHT filing'],
        ['August', 'VAT (July)', 'Aug 21', 'Previous month VAT filing'],
        ['August', 'PAYE (July)', 'Aug 10', 'Previous month PAYE remittance'],
        ['September', 'VAT (August)', 'Sep 21', 'Previous month VAT filing'],
        ['September', 'PAYE (August)', 'Sep 10', 'Previous month PAYE remittance'],
        ['October', 'VAT (September)', 'Oct 21', 'Previous month VAT filing'],
        ['October', 'PAYE (September)', 'Oct 10', 'Previous month PAYE remittance'],
        ['October', 'Q3 WHT Returns', 'Oct 21', 'Quarterly WHT filing'],
        ['November', 'VAT (October)', 'Nov 21', 'Previous month VAT filing'],
        ['November', 'PAYE (October)', 'Nov 10', 'Previous month PAYE remittance'],
        ['December', 'VAT (November)', 'Dec 21', 'Previous month VAT filing'],
        ['December', 'PAYE (November)', 'Dec 10', 'Previous month PAYE remittance'],
        ['December', 'CIT Monthly Installment', 'Dec 31', 'Final monthly payment for 2026'],
    ]

    for row_idx, row_data in enumerate(calendar_data, 2):
        for col_idx, cell_data in enumerate(row_data, 1):
            cell = ws4.cell(row=row_idx, column=col_idx, value=cell_data)
            cell.border = thin_border
            cell.alignment = left_align

    # Adjust column widths
    ws4.column_dimensions['A'].width = 15
    ws4.column_dimensions['B'].width = 30
    ws4.column_dimensions['C'].width = 15
    ws4.column_dimensions['D'].width = 40

    # -------------------------------------------------------------------------
    # SHEET 5: TAX RATES REFERENCE
    # -------------------------------------------------------------------------
    ws5 = wb.create_sheet('Tax Rates Reference')

    # Section 1: CIT Rates
    ws5.merge_cells('A1:C1')
    ws5.cell(row=1, column=1, value='COMPANY INCOME TAX RATES').font = header_font
    ws5.cell(row=1, column=1).fill = header_fill
    ws5.cell(row=1, column=1).alignment = center_align

    cit_rate_headers = ['Category', 'Criteria', 'Rate']
    for col, header in enumerate(cit_rate_headers, 1):
        cell = ws5.cell(row=2, column=col, value=header)
        cell.font = Font(bold=True)
        cell.border = thin_border

    cit_rates = [
        ['Small Company', 'Turnover ≤ NGN 50m AND Fixed Assets ≤ NGN 250m', '0%'],
        ['Medium Company', 'Turnover > NGN 50m to NGN 100m', '20%'],
        ['Large Company', 'Turnover > NGN 100m', '30%'],
    ]

    for row_idx, row_data in enumerate(cit_rates, 3):
        for col_idx, cell_data in enumerate(row_data, 1):
            cell = ws5.cell(row=row_idx, column=col_idx, value=cell_data)
            cell.border = thin_border

    # Section 2: PAYE Rates
    ws5.merge_cells('A7:C7')
    ws5.cell(row=7, column=1, value='PERSONAL INCOME TAX (PAYE) RATES').font = header_font
    ws5.cell(row=7, column=1).fill = header_fill
    ws5.cell(row=7, column=1).alignment = center_align

    paye_headers = ['Income Band (NGN)', 'Tax Rate', 'Cumulative Tax (NGN)']
    for col, header in enumerate(paye_headers, 1):
        cell = ws5.cell(row=8, column=col, value=header)
        cell.font = Font(bold=True)
        cell.border = thin_border

    paye_rates = [
        ['First 300,000', '7%', '21,000'],
        ['Next 300,000', '11%', '54,000'],
        ['Next 500,000', '15%', '129,000'],
        ['Next 500,000', '19%', '224,000'],
        ['Next 1,600,000', '21%', '560,000'],
        ['Above 3,200,000', '24%', 'Variable'],
    ]

    for row_idx, row_data in enumerate(paye_rates, 9):
        for col_idx, cell_data in enumerate(row_data, 1):
            cell = ws5.cell(row=row_idx, column=col_idx, value=cell_data)
            cell.border = thin_border

    # Section 3: WHT Rates
    ws5.merge_cells('A17:C17')
    ws5.cell(row=17, column=1, value='WITHHOLDING TAX RATES').font = header_font
    ws5.cell(row=17, column=1).fill = header_fill
    ws5.cell(row=17, column=1).alignment = center_align

    wht_headers = ['Payment Type', 'Companies', 'Individuals']
    for col, header in enumerate(wht_headers, 1):
        cell = ws5.cell(row=18, column=col, value=header)
        cell.font = Font(bold=True)
        cell.border = thin_border

    wht_rates = [
        ['Dividends', '10%', '10%'],
        ['Interest', '10%', '10%'],
        ['Rent', '10%', '10%'],
        ['Royalties', '10%', '10%'],
        ['Commission', '10%', '5%'],
        ['Consultancy/Professional Fees', '10%', '5%'],
        ['Technical Fees', '10%', '5%'],
        ['Construction', '5%', '5%'],
        ['Contracts (Government)', '5%', '5%'],
    ]

    for row_idx, row_data in enumerate(wht_rates, 19):
        for col_idx, cell_data in enumerate(row_data, 1):
            cell = ws5.cell(row=row_idx, column=col_idx, value=cell_data)
            cell.border = thin_border

    # Adjust column widths
    ws5.column_dimensions['A'].width = 35
    ws5.column_dimensions['B'].width = 25
    ws5.column_dimensions['C'].width = 25

    # Save workbook
    wb.save('/home/user/Kenna/Nigerian_Tax_Payment_Tracker_2026.xlsx')
    print("Excel tracker created successfully!")


# ============================================================================
# MAIN EXECUTION
# ============================================================================

if __name__ == '__main__':
    print("Creating Nigerian Tax Act 2025 documents...")
    print("-" * 50)

    create_word_document()
    create_excel_tracker()

    print("-" * 50)
    print("Documents created successfully!")
    print("\nFiles created:")
    print("1. Nigerian_Tax_Act_2025_Comprehensive_Review.docx")
    print("2. Nigerian_Tax_Payment_Tracker_2026.xlsx")
